"""Convert a copied ChatGPT Markdown response to Word and finish the task."""

from __future__ import annotations

import argparse
import re
import sys
from html.parser import HTMLParser
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from check_results import inspect_docx
from workflow_common import find_task, read_progress, task_docx_path, task_folder, utc_now, write_progress


HEADING_RE = re.compile(r"^(#{1,3})\s+(.+?)\s*$")
BULLET_RE = re.compile(r"^\s*[-*+]\s+(.+)$")
NUMBER_RE = re.compile(r"^\s*\d+[.)]\s+(.+)$")


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    heading_tokens = {
        "Heading 1": (16, "2E74B5", 16, 8),
        "Heading 2": (13, "2E74B5", 12, 6),
        "Heading 3": (12, "1F4D78", 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        style.paragraph_format.line_spacing = 1.167


def is_table_separator(line: str) -> bool:
    cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells)


def table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def add_markdown_table(document: Document, lines: list[str]) -> None:
    rows = [table_row(line) for line in lines if not is_table_separator(line)]
    if not rows:
        return
    width = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=width)
    table.style = "Table Grid"
    table.autofit = False
    column_width = Inches(6.5 / width)
    for row_index, values in enumerate(rows):
        for column_index in range(width):
            cell = table.cell(row_index, column_index)
            cell.width = column_width
            cell.text = values[column_index] if column_index < len(values) else ""
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
            if row_index == 0:
                set_cell_shading(cell, "F2F4F7")
                for run in cell.paragraphs[0].runs:
                    run.bold = True
    document.add_paragraph()


def markdown_to_docx(markdown: str, destination: Path) -> None:
    document = Document()
    configure_document(document)
    lines = markdown.replace("\r\n", "\n").split("\n")
    index = 0
    paragraph_lines: list[str] = []

    def flush_paragraph() -> None:
        if paragraph_lines:
            document.add_paragraph(" ".join(part.strip() for part in paragraph_lines))
            paragraph_lines.clear()

    while index < len(lines):
        line = lines[index].rstrip()
        heading = HEADING_RE.match(line)
        bullet = BULLET_RE.match(line)
        number = NUMBER_RE.match(line)
        if heading:
            flush_paragraph()
            document.add_heading(heading.group(2), level=len(heading.group(1)))
        elif bullet:
            flush_paragraph()
            document.add_paragraph(bullet.group(1), style="List Bullet")
        elif number:
            flush_paragraph()
            document.add_paragraph(number.group(1), style="List Number")
        elif "|" in line and index + 1 < len(lines) and is_table_separator(lines[index + 1]):
            flush_paragraph()
            table_lines = [line, lines[index + 1]]
            index += 2
            while index < len(lines) and "|" in lines[index] and lines[index].strip():
                table_lines.append(lines[index])
                index += 1
            add_markdown_table(document, table_lines)
            continue
        elif not line.strip():
            flush_paragraph()
        else:
            paragraph_lines.append(line)
        index += 1
    flush_paragraph()
    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(destination)


class WordHtmlParser(HTMLParser):
    """Convert the semantic subset emitted by ChatGPT into Word structures."""

    def __init__(self, document: Document):
        super().__init__(convert_charrefs=True)
        self.document = document
        self.paragraph = None
        self.list_stack: list[str] = []
        self.bold_depth = 0
        self.italic_depth = 0
        self.code_depth = 0
        self.pre_depth = 0
        self.blockquote_depth = 0
        self.table_rows: list[list[str]] | None = None
        self.table_row: list[str] | None = None
        self.table_cell: list[str] | None = None

    def _new_paragraph(self, style: str | None = None):
        self.paragraph = self.document.add_paragraph(style=style)
        if self.blockquote_depth:
            self.paragraph.paragraph_format.left_indent = Inches(0.35)
        return self.paragraph

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        tag = tag.lower()
        if tag in {"h1", "h2", "h3"}:
            self._new_paragraph(f"Heading {tag[1]}")
        elif tag == "p" and self.paragraph is None:
            self._new_paragraph()
        elif tag in {"ul", "ol"}:
            self.list_stack.append(tag)
        elif tag == "li":
            style = "List Number" if self.list_stack and self.list_stack[-1] == "ol" else "List Bullet"
            self._new_paragraph(style)
        elif tag in {"strong", "b"}:
            self.bold_depth += 1
        elif tag in {"em", "i"}:
            self.italic_depth += 1
        elif tag == "code":
            self.code_depth += 1
        elif tag == "pre":
            self.pre_depth += 1
            self._new_paragraph()
            self.paragraph.paragraph_format.left_indent = Inches(0.25)
        elif tag == "blockquote":
            self.blockquote_depth += 1
        elif tag == "br":
            if self.paragraph is None:
                self._new_paragraph()
            self.paragraph.add_run().add_break()
        elif tag == "table":
            self.table_rows = []
        elif tag == "tr" and self.table_rows is not None:
            self.table_row = []
        elif tag in {"th", "td"} and self.table_row is not None:
            self.table_cell = []

    def handle_endtag(self, tag: str) -> None:
        tag = tag.lower()
        if tag in {"h1", "h2", "h3", "p", "li", "pre"}:
            self.paragraph = None
        if tag in {"ul", "ol"} and self.list_stack:
            self.list_stack.pop()
        elif tag in {"strong", "b"}:
            self.bold_depth = max(0, self.bold_depth - 1)
        elif tag in {"em", "i"}:
            self.italic_depth = max(0, self.italic_depth - 1)
        elif tag == "code":
            self.code_depth = max(0, self.code_depth - 1)
        elif tag == "pre":
            self.pre_depth = max(0, self.pre_depth - 1)
        elif tag == "blockquote":
            self.blockquote_depth = max(0, self.blockquote_depth - 1)
        elif tag in {"th", "td"} and self.table_cell is not None and self.table_row is not None:
            self.table_row.append("".join(self.table_cell).strip())
            self.table_cell = None
        elif tag == "tr" and self.table_row is not None and self.table_rows is not None:
            if self.table_row:
                self.table_rows.append(self.table_row)
            self.table_row = None
        elif tag == "table" and self.table_rows is not None:
            self._add_table(self.table_rows)
            self.table_rows = None

    def handle_data(self, data: str) -> None:
        if self.table_cell is not None:
            self.table_cell.append(data)
            return
        if not self.pre_depth:
            data = re.sub(r"\s+", " ", data)
        if not data or (not data.strip() and self.paragraph is None):
            return
        if self.paragraph is None:
            self._new_paragraph()
        run = self.paragraph.add_run(data)
        run.bold = self.bold_depth > 0
        run.italic = self.italic_depth > 0
        if self.code_depth or self.pre_depth:
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            run.font.size = Pt(9.5)

    def _add_table(self, rows: list[list[str]]) -> None:
        if not rows:
            return
        width = max(len(row) for row in rows)
        table = self.document.add_table(rows=len(rows), cols=width)
        table.style = "Table Grid"
        table.autofit = False
        column_width = Inches(6.5 / width)
        for row_index, values in enumerate(rows):
            for column_index in range(width):
                cell = table.cell(row_index, column_index)
                cell.width = column_width
                cell.text = values[column_index] if column_index < len(values) else ""
                if row_index == 0:
                    set_cell_shading(cell, "F2F4F7")
                    for run in cell.paragraphs[0].runs:
                        run.bold = True
        self.document.add_paragraph()


def html_to_docx(html: str, destination: Path) -> None:
    document = Document()
    configure_document(document)
    parser = WordHtmlParser(document)
    parser.feed(html)
    parser.close()
    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(destination)


def main() -> int:
    parser = argparse.ArgumentParser(description="把复制的 GPT Markdown 内容转换成 Word 并完成任务")
    parser.add_argument("task_id", help="任务编号，例如 001")
    parser.add_argument("text_path", type=Path, help="复制内容保存的 UTF-8 Markdown 或文本文件")
    parser.add_argument("--chat-url", default="", help="可选：当前 ChatGPT 对话链接")
    parser.add_argument("--replace", action="store_true", help="允许替换已有 analysis.docx")
    parser.add_argument("--min-size-kb", type=int, default=20)
    parser.add_argument("--min-headings", type=int, default=3)
    parser.add_argument("--required-section", action="append", default=[])
    args = parser.parse_args()

    source = args.text_path.expanduser().resolve()
    if not source.is_file():
        print(f"复制内容文件不存在：{source}", file=sys.stderr)
        return 2
    try:
        markdown = source.read_text(encoding="utf-8-sig").strip()
    except UnicodeDecodeError:
        print("复制内容必须保存为 UTF-8 文本。", file=sys.stderr)
        return 2
    if not markdown:
        print("复制内容为空。", file=sys.stderr)
        return 2

    rows = read_progress()
    try:
        row = find_task(rows, args.task_id)
    except ValueError as exc:
        print(str(exc), file=sys.stderr)
        return 2
    folder = task_folder(row)
    destination = task_docx_path(row)
    if destination.exists() and not args.replace:
        print(f"目标已存在：{destination}。确认替换请添加 --replace。", file=sys.stderr)
        return 2

    candidate = folder / ".analysis.from_text.tmp.docx"
    markdown_to_docx(markdown, candidate)
    ok, message = inspect_docx(candidate, args.min_size_kb, args.min_headings, args.required_section)
    if destination.exists() and not ok:
        candidate.unlink(missing_ok=True)
        print(f"转换结果未通过检查，已保留原 Word：{message}", file=sys.stderr)
        return 1
    candidate.replace(destination)
    row["chat_url"] = args.chat_url or row["chat_url"]
    row["docx_status"] = "valid" if ok else "invalid"
    row["result_source"] = "copied_chat_text"
    row["quality_check"] = message
    row["status"] = "completed" if ok else "needs_review"
    row["error_message"] = "" if ok else message
    row["updated_at"] = utc_now()
    write_progress(rows)
    print(f"任务 {row['task_id']}：{'完成' if ok else '需要复核'} — {message}")
    return 0 if ok else 1


if __name__ == "__main__":
    raise SystemExit(main())
