"""Local queue bridge used by the personal ChatGPT Chrome extension."""

from __future__ import annotations

import argparse
from difflib import SequenceMatcher
import json
import os
import re
import secrets
import shutil
import subprocess
import threading
from http import HTTPStatus
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from urllib.parse import parse_qs, urlparse

from check_results import inspect_docx
from docx import Document
from workflow_common import OUTPUT_DIR, PROMPT_FILE, ROOT, adopt_single_task_docx, find_task, next_numeric_task_id, read_progress, safe_folder_component, task_docx_path, task_folder, task_pdf_path, utc_now, write_progress


HOST = "127.0.0.1"
DEFAULT_PORT = 8765
ROOT = Path(__file__).resolve().parents[1]
TOKEN_FILE = ROOT / ".workflow-token"
LOCK = threading.RLock()


def paste_clipboard_to_docx(destination: Path) -> None:
    """Let desktop Word perform the same Keep Source Formatting paste as a user."""
    destination = destination.resolve()
    destination.unlink(missing_ok=True)
    script = r"""
$ErrorActionPreference = 'Stop'
$word = $null
$document = $null
try {
    $word = New-Object -ComObject Word.Application
    $word.Visible = $false
    $word.DisplayAlerts = 0
    $document = $word.Documents.Add()
    # 16 = wdFormatOriginalFormatting
    $word.Selection.PasteAndFormat(16)
    # 16 = wdFormatDocumentDefault (.docx)
    $document.SaveAs2($env:WORKFLOW_DOCX_PATH, 16)
} finally {
    if ($null -ne $document) {
        $document.Close(0)
        [System.Runtime.InteropServices.Marshal]::FinalReleaseComObject($document) | Out-Null
    }
    if ($null -ne $word) {
        $word.Quit()
        [System.Runtime.InteropServices.Marshal]::FinalReleaseComObject($word) | Out-Null
    }
}
"""
    environment = os.environ.copy()
    environment["WORKFLOW_DOCX_PATH"] = str(destination)
    completed = subprocess.run(
        ["powershell.exe", "-NoProfile", "-STA", "-Command", script],
        capture_output=True,
        text=True,
        timeout=120,
        env=environment,
        check=False,
    )
    if completed.returncode != 0 or not destination.is_file():
        detail = (completed.stderr or completed.stdout or "Word 未生成文档").strip()
        raise RuntimeError(f"Word 保留源格式粘贴失败：{detail}")


def docx_matches_response(path: Path, response: str) -> bool:
    """Reject a Word paste if the clipboard changed after browser verification."""
    document = Document(path)
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        parts.extend(cell.text for row in table.rows for cell in row.cells)
    canonical = lambda value: re.sub(r"[\s#*_`>\-\[\]()|]+", "", value).casefold()
    expected = canonical(response)
    actual = canonical("\n".join(parts))
    if len(expected) < 80 or len(actual) < 80:
        return expected == actual
    if actual.startswith(expected[:120]) or expected.startswith(actual[:120]):
        return True
    return SequenceMatcher(None, expected[:10000], actual[:10000]).ratio() >= 0.75


def load_or_create_token() -> str:
    if TOKEN_FILE.exists():
        token = TOKEN_FILE.read_text(encoding="ascii").strip()
        if token:
            return token
    token = secrets.token_urlsafe(32)
    TOKEN_FILE.write_text(token + "\n", encoding="ascii")
    return token


class WorkflowHandler(BaseHTTPRequestHandler):
    server_version = "PaperWorkflow/1.0"

    @property
    def workflow_server(self) -> "WorkflowServer":
        return self.server  # type: ignore[return-value]

    def log_message(self, format: str, *args: object) -> None:
        print(f"[{self.log_date_time_string()}] {format % args}")

    def _cors_origin(self) -> str:
        origin = self.headers.get("Origin", "")
        if origin.startswith("chrome-extension://") or origin == "https://chatgpt.com":
            return origin
        return ""

    def _send_headers(self, status: int, content_type: str, length: int) -> None:
        self.send_response(status)
        origin = self._cors_origin()
        if origin:
            self.send_header("Access-Control-Allow-Origin", origin)
            self.send_header("Vary", "Origin")
        self.send_header("Access-Control-Allow-Headers", "Content-Type, X-Workflow-Token")
        self.send_header("Access-Control-Allow-Methods", "GET, POST, OPTIONS")
        self.send_header("Cache-Control", "no-store")
        self.send_header("Content-Type", content_type)
        self.send_header("Content-Length", str(length))
        self.end_headers()

    def _json(self, status: int, payload: object) -> None:
        data = json.dumps(payload, ensure_ascii=False).encode("utf-8")
        self._send_headers(status, "application/json; charset=utf-8", len(data))
        self.wfile.write(data)

    def _authorized(self, query: dict[str, list[str]]) -> bool:
        supplied = self.headers.get("X-Workflow-Token", "")
        if not supplied:
            supplied = query.get("token", [""])[0]
        return secrets.compare_digest(supplied, self.workflow_server.token)

    def _read_json(self) -> dict[str, object]:
        length = int(self.headers.get("Content-Length", "0"))
        if length <= 0 or length > 10 * 1024 * 1024:
            raise ValueError("请求正文为空或过大")
        value = json.loads(self.rfile.read(length).decode("utf-8"))
        if not isinstance(value, dict):
            raise ValueError("请求正文必须是 JSON 对象")
        return value

    def do_OPTIONS(self) -> None:  # noqa: N802
        self._send_headers(HTTPStatus.NO_CONTENT, "text/plain", 0)

    def do_GET(self) -> None:  # noqa: N802
        parsed = urlparse(self.path)
        query = parse_qs(parsed.query)
        if parsed.path == "/api/health":
            self._json(HTTPStatus.OK, {"ok": True, "service": "paper-workflow"})
            return
        if not self._authorized(query):
            self._json(HTTPStatus.UNAUTHORIZED, {"error": "invalid token"})
            return
        try:
            if parsed.path == "/api/state":
                self._state()
            elif parsed.path == "/api/next":
                self._next(query)
            elif parsed.path.startswith("/api/tasks/") and parsed.path.endswith("/pdf"):
                task_id = parsed.path.split("/")[3]
                self._pdf(task_id)
            else:
                self._json(HTTPStatus.NOT_FOUND, {"error": "not found"})
        except (ValueError, OSError) as exc:
            self._json(HTTPStatus.BAD_REQUEST, {"error": str(exc)})

    def do_POST(self) -> None:  # noqa: N802
        parsed = urlparse(self.path)
        query = parse_qs(parsed.query)
        if not self._authorized(query):
            self._json(HTTPStatus.UNAUTHORIZED, {"error": "invalid token"})
            return
        try:
            if parsed.path == "/api/import-pdf":
                self._import_pdf(query)
                return
            if parsed.path == "/api/open-output":
                self._open_output()
                return
            body = self._read_json()
            parts = parsed.path.strip("/").split("/")
            if len(parts) == 4 and parts[:2] == ["api", "tasks"] and parts[3] == "complete":
                self._complete(parts[2], body)
            elif len(parts) == 4 and parts[:2] == ["api", "tasks"] and parts[3] == "fail":
                self._fail(parts[2], body)
            elif len(parts) == 4 and parts[:2] == ["api", "tasks"] and parts[3] == "release":
                self._release(parts[2])
            elif parsed.path == "/api/status":
                self._set_status(body)
            elif parsed.path == "/api/delete-tasks":
                self._delete_tasks(body)
            else:
                self._json(HTTPStatus.NOT_FOUND, {"error": "not found"})
        except (ValueError, OSError, json.JSONDecodeError) as exc:
            self._json(HTTPStatus.BAD_REQUEST, {"error": str(exc)})

    def _open_output(self) -> None:
        OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
        subprocess.Popen(
            ["explorer.exe", str(OUTPUT_DIR.resolve())],
            close_fds=True,
        )
        self._json(HTTPStatus.OK, {"ok": True, "path": str(OUTPUT_DIR.resolve())})

    def _state(self) -> None:
        with LOCK:
            rows = read_progress()
            retained = []
            pruned = 0
            migrated = 0
            for row in rows:
                canonical_pdf = task_pdf_path(row)
                canonical_source = canonical_pdf.relative_to(ROOT).as_posix()
                if canonical_pdf.is_file() and row["source_path"] != canonical_source:
                    row["source_path"] = canonical_source
                    row["updated_at"] = utc_now()
                    migrated += 1
                source_exists = (ROOT / row["source_path"]).is_file()
                output_exists = task_folder(row).is_dir()
                if not source_exists and not output_exists:
                    pruned += 1
                else:
                    retained.append(row)
            if pruned or migrated:
                rows = retained
                write_progress(rows)
        counts: dict[str, int] = {}
        for row in rows:
            counts[row["status"]] = counts.get(row["status"], 0) + 1
        self._json(HTTPStatus.OK, {"counts": counts, "tasks": rows, "pruned": pruned, "migrated": migrated})

    def _import_pdf(self, query: dict[str, list[str]]) -> None:
        filename = query.get("filename", [""])[0].strip()
        filename = Path(filename).name
        if not filename or Path(filename).suffix.lower() != ".pdf":
            raise ValueError("只允许上传 PDF 文件")
        length = int(self.headers.get("Content-Length", "0"))
        if length <= 0 or length > 200 * 1024 * 1024:
            raise ValueError("PDF 为空或超过 200 MB")
        data = self.rfile.read(length)
        if len(data) != length or not data.startswith(b"%PDF-"):
            raise ValueError("文件内容不是有效的 PDF")
        with LOCK:
            rows = read_progress()
            if any(row["paper_name"].casefold() == filename.casefold() for row in rows):
                raise ValueError(f"任务列表中已存在同名论文：{filename}")
            task_id = f"{next_numeric_task_id(rows):03d}"
            paper_prefix = safe_folder_component(Path(filename).stem, max_length=30)
            row = {
                "task_id": task_id,
                "paper_name": filename,
                "source_path": "",
                "output_folder": (Path("output") / f"{task_id}_{paper_prefix}").as_posix(),
                "status": "pending",
                "chat_url": "",
                "docx_status": "missing",
                "result_source": "",
                "quality_check": "not_checked",
                "error_message": "",
                "updated_at": utc_now(),
            }
            destination = task_pdf_path(row)
            row["source_path"] = destination.relative_to(ROOT).as_posix()
            destination.parent.mkdir(parents=True, exist_ok=False)
            temporary = destination.parent / f".{destination.name}.uploading"
            try:
                temporary.write_bytes(data)
                temporary.replace(destination)
                rows.append(row)
                write_progress(rows)
            except BaseException:
                temporary.unlink(missing_ok=True)
                if destination.parent.exists():
                    shutil.rmtree(destination.parent, ignore_errors=True)
                raise
        self._json(HTTPStatus.OK, {"ok": True, "filename": filename, "size": length, "task_id": task_id})

    def _next(self, query: dict[str, list[str]]) -> None:
        requested = {
            item.strip().zfill(3)
            for value in query.get("ids", [])
            for item in value.split(",")
            if item.strip()
        }
        excluded = {
            item.strip().zfill(3)
            for value in query.get("exclude", [])
            for item in value.split(",")
            if item.strip()
        }
        with LOCK:
            rows = read_progress()
            row = next(
                (
                    item
                    for item in rows
                    if item["status"] in {"pending", "needs_review"}
                    and (not requested or item["task_id"] in requested)
                    and item["task_id"] not in excluded
                ),
                None,
            )
            if row is None:
                self._json(HTTPStatus.OK, {"done": True})
                return
            pdf = task_pdf_path(row)
            if not pdf.is_file() or not PROMPT_FILE.is_file():
                row["status"] = "failed"
                row["error_message"] = "缺少 original.pdf 或共享 analysis_prompt.md"
                row["updated_at"] = utc_now()
                write_progress(rows)
                raise ValueError(f"任务 {row['task_id']} 缺少 PDF 或共享提示词")
            prompt = PROMPT_FILE.read_text(encoding="utf-8").strip()
            if not prompt:
                raise ValueError("共享提示词 analysis_prompt.md 为空")
            row["status"] = "processing"
            row["error_message"] = ""
            row["updated_at"] = utc_now()
            write_progress(rows)
            payload = {
                "done": False,
                "task_id": row["task_id"],
                "paper_name": row["paper_name"],
                "request_text": prompt,
                "pdf_url": f"http://{HOST}:{self.workflow_server.server_port}/api/tasks/{row['task_id']}/pdf",
            }
        self._json(HTTPStatus.OK, payload)

    def _pdf(self, task_id: str) -> None:
        with LOCK:
            row = find_task(read_progress(), task_id)
            path = task_pdf_path(row)
        if not path.is_file():
            raise ValueError("PDF 不存在")
        data = path.read_bytes()
        self._send_headers(HTTPStatus.OK, "application/pdf", len(data))
        self.wfile.write(data)

    def _complete(self, task_id: str, body: dict[str, object]) -> None:
        response = str(body.get("response", "")).strip()
        response_source = str(body.get("response_source", "")).strip()
        clipboard_verified = body.get("clipboard_verified") is True
        chat_url = str(body.get("chat_url", "")).strip()
        if len(response) < 500:
            raise ValueError("网页回答过短，拒绝归档")
        if response_source != "chatgpt_selected_response" or not clipboard_verified:
            raise ValueError("只允许网页选中并经剪贴板校验的 ChatGPT 回答进入 Word；HTML 回退已关闭")
        with LOCK:
            rows = read_progress()
            row = find_task(rows, task_id)
            folder = task_pdf_path(row).parent
            folder.mkdir(parents=True, exist_ok=True)
            candidate = folder / ".analysis.extension.tmp.docx"
            try:
                paste_clipboard_to_docx(candidate)
            except (OSError, RuntimeError, subprocess.TimeoutExpired) as exc:
                candidate.unlink(missing_ok=True)
                raise ValueError(str(exc)) from exc
            if not docx_matches_response(candidate, response):
                candidate.unlink(missing_ok=True)
                raise ValueError("Word 粘贴内容与当前 GPT 回答不一致，已拒绝归档")
            used_source = "word_keep_source_format"
            # Word's Keep Source Formatting preserves visual headings but does
            # not reliably map them to built-in Word Heading styles.
            ok, message = inspect_docx(candidate, 20, 0, [])
            destination = task_docx_path(row)
            if destination.exists() and not ok:
                candidate.unlink(missing_ok=True)
            else:
                candidate.replace(destination)
            row["chat_url"] = chat_url
            row["docx_status"] = "valid" if ok else "invalid"
            row["result_source"] = used_source
            row["quality_check"] = message
            row["status"] = "completed" if ok else "needs_review"
            row["error_message"] = "" if ok else message
            row["updated_at"] = utc_now()
            write_progress(rows)
        self._json(HTTPStatus.OK, {"ok": ok, "message": message, "status": row["status"]})

    def _fail(self, task_id: str, body: dict[str, object]) -> None:
        message = str(body.get("error", "网页自动化失败")).strip()[:1000]
        chat_url = str(body.get("chat_url", "")).strip()
        with LOCK:
            rows = read_progress()
            row = find_task(rows, task_id)
            row["status"] = "needs_review"
            if chat_url:
                row["chat_url"] = chat_url
            row["error_message"] = message
            row["updated_at"] = utc_now()
            write_progress(rows)
        self._json(HTTPStatus.OK, {"ok": True})

    def _release(self, task_id: str) -> None:
        with LOCK:
            rows = read_progress()
            row = find_task(rows, task_id)
            if row["status"] == "processing":
                row["status"] = "pending"
                row["error_message"] = ""
                row["updated_at"] = utc_now()
                write_progress(rows)
        self._json(HTTPStatus.OK, {"ok": True, "status": row["status"]})

    def _set_status(self, body: dict[str, object]) -> None:
        raw_ids = body.get("task_ids", [])
        if not isinstance(raw_ids, list):
            raise ValueError("task_ids 必须是列表")
        task_ids = {str(value).strip().zfill(3) for value in raw_ids if str(value).strip()}
        status = str(body.get("status", "")).strip()
        allowed = {"pending", "needs_review", "completed", "failed"}
        if not task_ids:
            raise ValueError("没有选择任务")
        if status not in allowed:
            raise ValueError("不允许手动设置该状态")
        with LOCK:
            rows = read_progress()
            found: set[str] = set()
            for row in rows:
                if row["task_id"] not in task_ids:
                    continue
                found.add(row["task_id"])
                if status == "completed":
                    docx_path = adopt_single_task_docx(row)
                    ok, message = inspect_docx(docx_path, 20, 0, [])
                    if not ok:
                        raise ValueError(f"任务 {row['task_id']} 不能设为 completed：{message}")
                    row["docx_status"] = "valid"
                    row["quality_check"] = message
                    row["error_message"] = ""
                elif status == "pending":
                    row["error_message"] = ""
                elif status == "needs_review":
                    row["error_message"] = "用户手动标记为需要复核"
                elif status == "failed":
                    row["error_message"] = "用户手动标记为失败"
                row["status"] = status
                row["updated_at"] = utc_now()
            missing = sorted(task_ids - found)
            if missing:
                raise ValueError("找不到任务：" + ", ".join(missing))
            write_progress(rows)
        self._json(HTTPStatus.OK, {"ok": True, "updated": sorted(task_ids), "status": status})

    def _delete_tasks(self, body: dict[str, object]) -> None:
        raw_ids = body.get("task_ids", [])
        if not isinstance(raw_ids, list):
            raise ValueError("task_ids 必须是列表")
        task_ids = {str(value).strip().zfill(3) for value in raw_ids if str(value).strip()}
        if not task_ids:
            raise ValueError("没有选择任务")
        with LOCK:
            rows = read_progress()
            by_id = {row["task_id"]: row for row in rows}
            missing = sorted(task_ids - by_id.keys())
            if missing:
                raise ValueError("找不到任务：" + ", ".join(missing))
            deleted: list[str] = []
            failures: dict[str, str] = {}
            for task_id in sorted(task_ids):
                row = by_id[task_id]
                try:
                    output = task_folder(row).resolve()
                    output.relative_to(OUTPUT_DIR.resolve())
                    if output.exists():
                        shutil.rmtree(output)
                    deleted.append(task_id)
                except (OSError, ValueError) as exc:
                    failures[task_id] = str(exc)
            if deleted:
                rows = [row for row in rows if row["task_id"] not in set(deleted)]
                write_progress(rows)
        self._json(HTTPStatus.OK, {"ok": True, "all_deleted": not failures, "deleted": deleted, "failures": failures})


class WorkflowServer(ThreadingHTTPServer):
    def __init__(self, address: tuple[str, int], token: str):
        self.token = token
        super().__init__(address, WorkflowHandler)


def main() -> int:
    parser = argparse.ArgumentParser(description="为个人 Chrome 扩展提供本地论文任务队列")
    parser.add_argument("--port", type=int, default=DEFAULT_PORT)
    args = parser.parse_args()
    token = load_or_create_token()
    server = WorkflowServer((HOST, args.port), token)
    print(f"本地服务：http://{HOST}:{args.port}")
    print(f"扩展令牌：{token}")
    print("仅监听本机；按 Ctrl+C 停止。")
    try:
        server.serve_forever()
    except KeyboardInterrupt:
        print("\n服务已停止。")
    finally:
        server.server_close()
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
