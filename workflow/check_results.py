"""Perform a lightweight local structural check of generated Word files."""

from __future__ import annotations

import argparse
import sys
from pathlib import Path

from workflow_common import find_task, read_progress, task_docx_path, utc_now, write_progress


def inspect_docx(docx_path: Path, min_size_kb: int, min_headings: int, required_sections: list[str]) -> tuple[bool, str]:
    if not docx_path.exists():
        return False, "analysis.docx 不存在"
    if docx_path.stat().st_size < min_size_kb * 1024:
        return False, f"文件小于 {min_size_kb} KB"
    try:
        from docx import Document
    except ImportError:
        return False, "缺少 python-docx；请运行 py -m pip install -r requirements.txt"

    try:
        document = Document(docx_path)
    except Exception as exc:  # python-docx raises several exception types for corrupt files
        return False, f"无法打开 Word：{exc}"

    headings = [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text.strip() and paragraph.style and paragraph.style.name.lower().startswith("heading")
    ]
    table_text = [
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    ]
    text = "\n".join(
        [*(paragraph.text for paragraph in document.paragraphs), *table_text]
    ).casefold()
    missing = [section for section in required_sections if section.casefold() not in text]
    if len(headings) < min_headings:
        return False, f"仅检测到 {len(headings)} 个标题，要求至少 {min_headings} 个"
    if missing:
        return False, "缺少指定章节关键词：" + "、".join(missing)
    return True, f"可打开；{docx_path.stat().st_size // 1024} KB；检测到 {len(headings)} 个标题"


def main() -> int:
    parser = argparse.ArgumentParser(description="检查一个或全部任务的 analysis.docx")
    parser.add_argument("--task-id", help="只检查指定任务，例如 001")
    parser.add_argument("--min-size-kb", type=int, default=20)
    parser.add_argument("--min-headings", type=int, default=3)
    parser.add_argument("--required-section", action="append", default=[], help="要求出现的章节关键词；可重复使用")
    args = parser.parse_args()

    rows = read_progress()
    if not rows:
        print("progress.csv 不存在或为空。请先通过浏览器扩展导入 PDF。", file=sys.stderr)
        return 2
    if args.task_id:
        selected = [find_task(rows, args.task_id)]
    else:
        selected = [row for row in rows if task_docx_path(row).is_file()]
        if not selected:
            print("没有已归档的 analysis.docx 可检查。")
            return 0
    failures = 0
    for row in selected:
        ok, message = inspect_docx(
            task_docx_path(row),
            args.min_size_kb,
            args.min_headings,
            args.required_section,
        )
        row["docx_status"] = "valid" if ok else "invalid"
        row["quality_check"] = message
        row["updated_at"] = utc_now()
        if ok:
            row["status"] = "completed"
            row["error_message"] = ""
            print(f"{row['task_id']}: 通过 — {message}")
        else:
            row["status"] = "needs_review"
            row["error_message"] = message
            failures += 1
            print(f"{row['task_id']}: 需检查 — {message}")
    write_progress(rows)
    return 1 if failures else 0


if __name__ == "__main__":
    raise SystemExit(main())
